"""
Standalone image processor that works independently of the existing placeholder system.
Uses docxtpl to handle image insertion in table cells.
"""

import os
import tempfile
import shutil
from typing import Optional
from pathlib import Path
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Inches
import io
import serpapi
import requests


def add_images_to_document(
    input_docx_path: str,
    output_docx_path: str,
    entity_name: str,
    serp_api_key: str
) -> bool:
    """
    Standalone function to add images to a DOCX document.
    
    This function:
    1. Takes an existing DOCX file (with or without other placeholders already filled)
    2. Searches for [[IMAGE]] placeholders
    3. Converts them to docxtpl format and inserts actual images
    4. Saves the result to a new file
    
    Args:
        input_docx_path: Path to the input DOCX file
        output_docx_path: Path where to save the output DOCX file
        entity_name: Name of the entity to search images for
        serp_api_key: SerpAPI key for Google Images search
        
    Returns:
        True if successful, False otherwise
    """
    try:
        # Step 1: Check if document has [[IMAGE]] placeholders
        if not _has_image_placeholders(input_docx_path):
            raise RuntimeError(f"No [[IMAGE]] placeholders found in template for {entity_name}")
        
        # Step 2: Convert [[IMAGE]] to {{image_placeholder}} in a temp file
        temp_template_path = _convert_image_placeholders_to_docxtpl_format(input_docx_path)
        
        # Step 3: Search and download image
        image_bytes = _search_and_download_image(entity_name, serp_api_key)
        
        # Step 4: Process with docxtpl
        doc = DocxTemplate(temp_template_path)
        
        if image_bytes:
            # Validate and convert image format for compatibility
            try:
                from PIL import Image as PILImage
                
                image_bytes.seek(0)
                pil_image = PILImage.open(image_bytes)
                
                # Convert to RGB if needed (handles grayscale, transparency, etc.)
                if pil_image.mode in ('RGBA', 'P', 'LA', 'L'):
                    if pil_image.mode == 'L':
                        # Grayscale to RGB
                        pil_image = pil_image.convert('RGB')
                    elif pil_image.mode in ('RGBA', 'LA'):
                        # Handle transparency with white background
                        rgb_image = PILImage.new('RGB', pil_image.size, (255, 255, 255))
                        rgb_image.paste(pil_image, mask=pil_image.split()[-1])
                        pil_image = rgb_image
                    elif pil_image.mode == 'P':
                        # Palette mode
                        pil_image = pil_image.convert('RGBA')
                        rgb_image = PILImage.new('RGB', pil_image.size, (255, 255, 255))
                        rgb_image.paste(pil_image, mask=pil_image.split()[-1])
                        pil_image = rgb_image
                    elif pil_image.mode != 'RGB':
                        pil_image = pil_image.convert('RGB')
                
                # Save as PNG for compatibility
                converted_bytes = io.BytesIO()
                pil_image.save(converted_bytes, format='PNG', optimize=True)
                converted_bytes.seek(0)
                
                # Create InlineImage object
                image_obj = InlineImage(doc, converted_bytes, width=Inches(2.0))
                context = {'image_placeholder': image_obj}
                
            except Exception as img_error:
                print(f"Image conversion failed: {img_error}")
                context = {'image_placeholder': ''}
        else:
            # No image available - remove placeholder
            context = {'image_placeholder': ''}
        
        # Step 5: Render and save
        doc.render(context)
        doc.save(output_docx_path)
        
        # Step 6: Cleanup temp file
        try:
            os.remove(temp_template_path)
        except:
            pass
        
        return True
        
    except Exception as e:
        print(f"Failed to process images: {e}")
        
        # Fallback: copy input to output if processing fails
        try:
            shutil.copy2(input_docx_path, output_docx_path)
        except:
            pass
            
        return False


def _has_image_placeholders(docx_path: str) -> bool:
    """Check if document contains [[IMAGE]] placeholders."""
    try:
        from docx import Document
        from docx.text.paragraph import Paragraph
        from docx.oxml.ns import qn
        import re
        
        document = Document(docx_path)
        pattern = re.compile(r"\[\[IMAGE\]\]")
        
        def _check_element_for_image_placeholders(element):
            """Check an XML element for IMAGE placeholders."""
            for para_element in element.iter(qn("w:p")):
                para = Paragraph(para_element, para_element.getparent())
                full_text = "".join(run.text for run in para.runs)
                if pattern.search(full_text):
                    return True
            return False
        
        # Check main document body
        if _check_element_for_image_placeholders(document._element.body):
            return True
        
        # Check headers and footers
        for section in document.sections:
            if _check_element_for_image_placeholders(section.header._element):
                return True
            if _check_element_for_image_placeholders(section.footer._element):
                return True
        
        return False
        
    except Exception:
        return False


def _convert_image_placeholders_to_docxtpl_format(input_path: str) -> str:
    """Create a temporary copy with [[IMAGE]] converted to {{image_placeholder}}."""
    from docx import Document
    from docx.text.paragraph import Paragraph
    from docx.oxml.ns import qn
    import re
    
    # Create temporary file
    temp_fd, temp_path = tempfile.mkstemp(suffix='.docx')
    os.close(temp_fd)
    
    # Copy and modify
    shutil.copy2(input_path, temp_path)
    document = Document(temp_path)
    
    pattern = re.compile(r"\[\[IMAGE\]\]")
    
    def _replace_in_element(element):
        """Replace IMAGE placeholders in an XML element."""
        for para_element in element.iter(qn("w:p")):
            para = Paragraph(para_element, para_element.getparent())
            
            # Check for placeholders across all runs in the paragraph
            full_text = "".join(run.text for run in para.runs)
            if not pattern.search(full_text):
                continue
            
            # Use re.split to handle the replacement properly
            parts = pattern.split(full_text)
            
            # Clear the paragraph content and rebuild it
            p = para._p
            p.clear_content()

            for i, part in enumerate(parts):
                if i % 2 == 1:  # This would be the IMAGE match position
                    # Add the docxtpl placeholder
                    para.add_run("{{image_placeholder}}")
                else:  # This is regular text
                    if part:
                        para.add_run(part)
    
    # Replace in main document body
    _replace_in_element(document._element.body)
    
    # Replace in headers and footers
    for section in document.sections:
        _replace_in_element(section.header._element)
        _replace_in_element(section.footer._element)
    
    # Save modified document
    document.save(temp_path)
    
    return temp_path


def _search_and_download_image(entity_name: str, serp_api_key: str) -> Optional[io.BytesIO]:
    """Search for and download an image for the given entity."""
    try:
        # Search for image using SerpAPI
        client = serpapi.Client(api_key=serp_api_key)
        params = {
            "engine": "google_images",
            "q": f"{entity_name} portrait",
            "tbm": "isch",
            "num": 15,
            "safe": "active"
        }
        
        results = client.search(params)
        images = results.get("images_results", [])
        
        if not images:
            return None
        
        # Try to download images
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
        max_images = min(15, len(images))
        
        for i in range(max_images):
            image = images[i]
            image_url = image.get("original")
            if not image_url:
                continue
            
            try:
                response = requests.get(image_url, headers=headers, timeout=10)
                response.raise_for_status()
                
                # Check if it's actually an image
                content_type = response.headers.get('Content-Type', '')
                if 'image' in content_type:
                    return io.BytesIO(response.content)
                    
            except Exception:
                continue
        
        return None
        
    except Exception:
        return None


# Convenience function for easy integration
def process_document_images(
    document_path: str,
    entity_name: str, 
    serp_api_key: str,
    output_suffix: str = "_with_images"
) -> str:
    """
    Convenience function that processes images and returns the new file path.
    
    Args:
        document_path: Path to the document
        entity_name: Name for image search
        serp_api_key: SerpAPI key
        output_suffix: Suffix to add to filename
        
    Returns:
        Path to the new document with images
    """
    # Generate output path
    path_obj = Path(document_path)
    output_path = str(path_obj.parent / f"{path_obj.stem}{output_suffix}{path_obj.suffix}")
    
    # Process images
    success = add_images_to_document(
        input_docx_path=document_path,
        output_docx_path=output_path,
        entity_name=entity_name,
        serp_api_key=serp_api_key
    )
    
    return output_path if success else document_path
