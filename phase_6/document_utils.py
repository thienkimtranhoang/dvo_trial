"""Document manipulation utilities for DOCX files."""

import re
from typing import Dict, List, Tuple
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Local constant for unknown dates
UNKNOWN_PUBLICATION_DATE = "Unknown Publication Date"


def convert_merged_facts_to_flat_format(entries: List[Dict]) -> List[Tuple[str, List[str]]]:
    """
    Convert LLM-merged fact entries to flattened (value, sources) format for document generation.
    
    Args:
        entries: List of fact entries in various formats
        
    Returns:
        List of (value, sources) tuples ready for document insertion
    """
    if not entries:
        return []

    output = []
    for entry in entries:
        # Handle the LLM-merged format: {"value": "...", "sources": ["src1", "src2"]}
        if isinstance(entry, dict) and "sources" in entry:
            # LLM-merged format (preferred)
            value = str(entry["value"]).strip()
            sources = sorted(list(set(entry["sources"])))  # Remove duplicates and sort
            output.append((value, sources))
            
        elif isinstance(entry, dict) and "value" in entry:
            # Fallback: Original format {"value": "...", "source": "..."}
            value = str(entry["value"]).strip()
            sources = [entry.get("source", UNKNOWN_PUBLICATION_DATE)]
            output.append((value, sources))
            
        else:
            # Handle any unexpected format
            value = str(entry).strip()
            sources = [UNKNOWN_PUBLICATION_DATE]
            output.append((value, sources))
    
    return output

def convert_merged_facts_with_dates_to_flat_format(entries: List[Dict]) -> List[Tuple[str, List[Tuple[str, str]]]]:
    """
    Convert LLM-merged fact entries with dates to flattened format for document generation.
    
    Args:
        entries: List of fact entries with date information
        Format: [{"value": "...", "sources": [{"link": "...", "date": "..."}]}]
        
    Returns:
        List of (value, [(url, date), ...]) tuples ready for document insertion
    """
    if not entries:
        return []

    output = []
    for entry in entries:
        if isinstance(entry, dict) and "sources" in entry:
            value = str(entry["value"]).strip()
            
            # Handle new format with date objects
            source_tuples = []
            for source in entry["sources"]:
                if isinstance(source, dict) and "link" in source and "date" in source:
                    # New format: {"link": "url", "date": "dd-mm-yyyy"}
                    url = source["link"]
                    date = source["date"]
                    source_tuples.append((url, date))
                else:
                    # Fallback for simple string sources
                    url = str(source) if not isinstance(source, dict) else source.get("link", str(source))
                    source_tuples.append((url, "Unknown Publication Date"))
            
            output.append((value, source_tuples))
        else:
            # Handle unexpected format
            value = str(entry).strip()
            source_tuples = [("Unknown Source", "Unknown Publication Date")]
            output.append((value, source_tuples))
    
    return output


def populate_document_placeholders(document: Document, field_data: Dict[str, List[Tuple[str, List[str]]]]) -> None:
    """
    Replace placeholders in a DOCX document with actual data and source links.
    
    Args:
        document: The DOCX document to populate
        field_data: Dictionary mapping field names to (value, sources) tuples
    """
    pattern = re.compile(r"\[\[([A-Za-z0-9_]+)\]\]")
    
    # Replace in main document body
    _replace_placeholders_in_element(document._element.body, field_data, pattern, document)

    # Replace in headers and footers
    for section in document.sections:
        _replace_placeholders_in_element(section.header._element, field_data, pattern, document)
        _replace_placeholders_in_element(section.footer._element, field_data, pattern, document)

def populate_document_placeholders_with_dates(document: Document, field_data: Dict[str, List[Tuple[str, List[Tuple[str, str]]]]]) -> None:
    """
    Replace placeholders in a DOCX document with actual data and source links WITH DATES.
    
    Args:
        document: The DOCX document to populate
        field_data: Dictionary mapping field names to (value, [(url, date), ...]) tuples
    """
    pattern = re.compile(r"\[\[([A-Za-z0-9_]+)\]\]")
    
    # Replace in main document body
    _replace_placeholders_with_dates_in_element(document._element.body, field_data, pattern, document)

    # Replace in headers and footers
    for section in document.sections:
        _replace_placeholders_with_dates_in_element(section.header._element, field_data, pattern, document)
        _replace_placeholders_with_dates_in_element(section.footer._element, field_data, pattern, document)


def _replace_placeholders_in_element(element, replacements: Dict, pattern: re.Pattern, document: Document) -> None:
    """
    Replace placeholders in a specific document element.
    
    Args:
        element: XML element to process
        replacements: Dictionary of field replacements
        pattern: Compiled regex pattern for finding placeholders
        document: Parent document for hyperlink creation
    """
    source_counter = 1
    
    for para_element in element.iter(qn("w:p")):
        para = Paragraph(para_element, para_element.getparent())
        
        # Check for placeholders across all runs in the paragraph
        full_text = "".join(run.text for run in para.runs)
        if not pattern.search(full_text):
            continue

        # Use re.split to handle multiple placeholders and surrounding text
        parts = pattern.split(full_text)
        
        # Clear the paragraph content and rebuild it
        p = para._p
        p.clear_content()

        for i, part in enumerate(parts):
            if i % 2 == 1:  # This is a placeholder key
                key = part
                val_list = replacements.get(key)

                if val_list:
                    for item_idx, (text, sources) in enumerate(val_list):
                        if item_idx > 0:
                            para.add_run("\n\n")
                        para.add_run(text + " ")
                        for src in sources:
                            label = f"[Source {source_counter}]"
                            add_hyperlink_to_paragraph(para, src, label, document)
                            para.add_run(" ")
                            source_counter += 1
            else:  # This is a static text part
                if part:
                    para.add_run(part)

def _replace_placeholders_with_dates_in_element(element, replacements: Dict, pattern: re.Pattern, document: Document) -> None:
    """
    Replace placeholders in a specific document element WITH DATE INFORMATION.
    
    Args:
        element: XML element to process
        replacements: Dictionary of field replacements with date info
        pattern: Compiled regex pattern for finding placeholders
        document: Parent document for hyperlink creation
    """
    source_counter = 1
    
    for para_element in element.iter(qn("w:p")):
        para = Paragraph(para_element, para_element.getparent())
        
        # Check for placeholders across all runs in the paragraph
        full_text = "".join(run.text for run in para.runs)
        if not pattern.search(full_text):
            continue

        # Use re.split to handle multiple placeholders and surrounding text
        parts = pattern.split(full_text)
        
        # Clear the paragraph content and rebuild it
        p = para._p
        p.clear_content()

        for i, part in enumerate(parts):
            if i % 2 == 1:  # This is a placeholder key
                key = part
                val_list = replacements.get(key)

                if val_list:
                    for item_idx, (text, source_tuples) in enumerate(val_list):
                        if item_idx > 0:
                            para.add_run("\n\n")

                        # Replace [[N]] and [[N][M]] inline citations with hyperlinks
                        # Handle both [[1]] and [[1][2][3]] patterns
                        # Handle both [[N]] numeric citations and [[source]] named citations
                        inline_pattern = re.compile(r"\[\[(source|\d+(?:(?:\]\[|\],\s*\[)\d+)*)\]\]")
                        parts_inner = inline_pattern.split(text)
                        for j, part_inner in enumerate(parts_inner):
                            if j % 2 == 1:
                                if part_inner == "source":
                                    # [[source]] — use first source URL
                                    if source_tuples:
                                        url, date = source_tuples[0]
                                        if url:
                                            add_hyperlink_to_paragraph(para, url, "[source]", document)
                                        else:
                                            para.add_run("[source]")
                                else:
                                    # Numeric citations [[1]] or [[1][2]]
                                    nums = [int(x) for x in re.findall(r"\d+", part_inner)]
                                    for n in nums:
                                        if 1 <= n <= len(source_tuples):
                                            url, date = source_tuples[n - 1]
                                            if url:
                                                add_hyperlink_to_paragraph(para, url, f"[{n}]", document)
                                            else:
                                                para.add_run(f"[{n}]")
                                        else:
                                            para.add_run(f"[{n}]")
                            else:
                                if part_inner:
                                    para.add_run(part_inner)
            else:  # This is a static text part
                if part:
                    para.add_run(part)


def add_hyperlink_to_paragraph(paragraph: Paragraph, url: str, text: str, document: Document, 
                              color: str = "0000FF", underline: bool = True) -> None:
    """
    Add a hyperlink to a paragraph in a DOCX document.
    
    Args:
        paragraph: The paragraph to add the hyperlink to
        url: The URL for the hyperlink
        text: The display text for the hyperlink
        document: The parent document
        color: Hex color code for the hyperlink (default: blue)
        underline: Whether to underline the hyperlink
        
    Raises:
        TypeError: If paragraph is not a Paragraph object
    """
    if not isinstance(paragraph, Paragraph):
        raise TypeError("Expected a docx.text.paragraph.Paragraph object.")

    part = document.part
    r_id = part.relate_to(
        url,
        reltype="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    # Create the w:hyperlink tag
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    # Create a new w:r (run)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Set color
    color_elem = OxmlElement("w:color")
    color_elem.set(qn("w:val"), color)
    rPr.append(color_elem)

    # Set underline
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single" if underline else "none")
    rPr.append(u)

    new_run.append(rPr)

    # Add text content
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)

    # Append the hyperlink to the paragraph's XML element
    paragraph._p.append(hyperlink)
